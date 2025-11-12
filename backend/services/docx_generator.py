"""
DOCX Resume Generator Service
Generates Word documents (.docx) from resume data
Related to UC-051: Resume Export and Formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, List, Any


class DOCXGenerator:
    """Generate DOCX documents from resume data"""

    # Color hex to RGB conversion
    @staticmethod
    def hex_to_rgb(hex_color: str) -> tuple:
        """Convert hex color string to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    @staticmethod
    def generate_docx_from_resume(resume_data: dict) -> bytes:
        """
        Generate DOCX document from resume data

        Args:
            resume_data: Resume object from MongoDB

        Returns:
            DOCX file content as bytes
        """
        doc = Document()

        # Set margins (1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Get template colors for styling
        colors = resume_data.get('colors', {})
        primary_color = colors.get('primary', '#000000')
        primary_rgb = DOCXGenerator.hex_to_rgb(primary_color)

        # Add header with name and contact
        DOCXGenerator._add_header(doc, resume_data, primary_rgb)

        # Add summary if present
        if resume_data.get('summary'):
            DOCXGenerator._add_section(doc, 'Professional Summary', resume_data['summary'], primary_rgb)

        # Add experience
        if resume_data.get('experience'):
            DOCXGenerator._add_experience_section(doc, resume_data['experience'], primary_rgb)

        # Add education
        if resume_data.get('education'):
            DOCXGenerator._add_education_section(doc, resume_data['education'], primary_rgb)

        # Add skills
        if resume_data.get('skills'):
            DOCXGenerator._add_skills_section(doc, resume_data['skills'], primary_rgb)

        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io.getvalue()

    @staticmethod
    def _add_header(doc: Document, resume_data: dict, primary_rgb: tuple) -> None:
        """Add header with name and contact information"""
        contact = resume_data.get('contact', {})

        # Add name
        name = contact.get('name', 'Resume')
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(*primary_rgb)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add contact information
        contact_info = []
        if contact.get('email'):
            contact_info.append(contact['email'])
        if contact.get('phone'):
            contact_info.append(contact['phone'])
        if contact.get('address'):
            contact_info.append(contact['address'])
        if contact.get('linkedin'):
            contact_info.append(contact['linkedin'])

        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(10)

        # Add space after header
        doc.add_paragraph()

    @staticmethod
    def _add_section(doc: Document, title: str, content: str, primary_rgb: tuple) -> None:
        """Add a simple section with title and content"""
        # Add section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*primary_rgb)
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(3)

        # Add separator line
        separator = doc.add_paragraph('_' * 80)
        separator_format = separator.paragraph_format
        separator_format.space_before = Pt(0)
        separator_format.space_after = Pt(6)
        for run in separator.runs:
            run.font.color.rgb = RGBColor(*primary_rgb)
            run.font.size = Pt(10)

        # Add content
        content_para = doc.add_paragraph(content)
        content_para.paragraph_format.space_after = Pt(12)

    @staticmethod
    def _add_experience_section(doc: Document, experience: List[Dict[str, Any]], primary_rgb: tuple) -> None:
        """Add work experience section"""
        # Add section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('Professional Experience')
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*primary_rgb)
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(3)

        # Add separator line
        separator = doc.add_paragraph('_' * 80)
        separator_format = separator.paragraph_format
        separator_format.space_before = Pt(0)
        separator_format.space_after = Pt(6)
        for run in separator.runs:
            run.font.color.rgb = RGBColor(*primary_rgb)
            run.font.size = Pt(10)

        for idx, exp in enumerate(experience):
            # Job title and company
            header_para = doc.add_paragraph()
            title_run = header_para.add_run(exp.get('title', 'Position'))
            title_run.font.bold = True
            title_run.font.size = Pt(11)

            if exp.get('company'):
                header_para.add_run(f" at {exp['company']}")

            # Dates
            dates = []
            if exp.get('start_date'):
                dates.append(exp['start_date'])
            if exp.get('end_date'):
                dates.append(exp['end_date'])
            if dates:
                date_str = ' - '.join(dates)
                date_para = doc.add_paragraph(date_str)
                date_para.paragraph_format.space_before = Pt(0)
                date_para.paragraph_format.space_after = Pt(3)
                for run in date_para.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)

            # Location
            if exp.get('location'):
                location_para = doc.add_paragraph(exp['location'])
                location_para.paragraph_format.space_before = Pt(0)
                location_para.paragraph_format.space_after = Pt(6)
                for run in location_para.runs:
                    run.font.size = Pt(10)

            # Description
            if exp.get('description'):
                desc_para = doc.add_paragraph(exp['description'], style='List Bullet')
                desc_para.paragraph_format.space_after = Pt(6)

            # Space between entries
            if idx < len(experience) - 1:
                doc.add_paragraph()

    @staticmethod
    def _add_education_section(doc: Document, education: List[Dict[str, Any]], primary_rgb: tuple) -> None:
        """Add education section"""
        # Add section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('Education')
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*primary_rgb)
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(3)

        # Add separator line
        separator = doc.add_paragraph('_' * 80)
        separator_format = separator.paragraph_format
        separator_format.space_before = Pt(0)
        separator_format.space_after = Pt(6)
        for run in separator.runs:
            run.font.color.rgb = RGBColor(*primary_rgb)
            run.font.size = Pt(10)

        for idx, edu in enumerate(education):
            # Degree and school
            header_para = doc.add_paragraph()
            degree_run = header_para.add_run(edu.get('degree', 'Degree'))
            degree_run.font.bold = True
            degree_run.font.size = Pt(11)

            if edu.get('school'):
                header_para.add_run(f" from {edu['school']}")

            # Graduation year
            if edu.get('graduation_year'):
                year_para = doc.add_paragraph(f"Graduated: {edu['graduation_year']}")
                year_para.paragraph_format.space_before = Pt(0)
                year_para.paragraph_format.space_after = Pt(6)
                for run in year_para.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)

            # GPA
            if edu.get('gpa'):
                gpa_para = doc.add_paragraph(f"GPA: {edu['gpa']}")
                gpa_para.paragraph_format.space_before = Pt(0)
                gpa_para.paragraph_format.space_after = Pt(6)
                for run in gpa_para.runs:
                    run.font.size = Pt(10)

            # Description
            if edu.get('description'):
                desc_para = doc.add_paragraph(edu['description'], style='List Bullet')
                desc_para.paragraph_format.space_after = Pt(6)

            # Space between entries
            if idx < len(education) - 1:
                doc.add_paragraph()

    @staticmethod
    def _add_skills_section(doc: Document, skills: List[str], primary_rgb: tuple) -> None:
        """Add skills section"""
        # Add section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('Skills')
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*primary_rgb)
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(3)

        # Add separator line
        separator = doc.add_paragraph('_' * 80)
        separator_format = separator.paragraph_format
        separator_format.space_before = Pt(0)
        separator_format.space_after = Pt(6)
        for run in separator.runs:
            run.font.color.rgb = RGBColor(*primary_rgb)
            run.font.size = Pt(10)

        # Add skills as bullet points
        for skill in skills:
            skill_text = skill if isinstance(skill, str) else skill.get('name', '')
            if skill_text:
                skill_para = doc.add_paragraph(skill_text, style='List Bullet')
                skill_para.paragraph_format.space_after = Pt(3)
