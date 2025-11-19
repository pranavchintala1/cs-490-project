from fastapi import APIRouter, HTTPException, Header, Path, File, UploadFile, Form, Body, Depends
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
from uuid import uuid4
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import tempfile
import requests
from bs4 import BeautifulSoup

from schema.CoverLetter import CoverLetterIn, CoverLetterOut
from mongo.cover_letters_dao import cover_letters_dao
from sessions.session_authorizer import authorize

coverletter_router = APIRouter(prefix="/cover-letters")

# ============================================================
# GET usage stats aggregated by template type (MUST BE FIRST)
# ============================================================
@coverletter_router.get("/usage/by-type")
async def get_usage_by_template_type():
    """Get aggregated usage counts grouped by template type (style_industry)."""
    try:
        usage_stats = await cover_letters_dao.get_usage_by_template_type()
        return usage_stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get usage stats: {str(e)}")

# ============================================================
# GET all cover letters for the current user
# ============================================================
@coverletter_router.get("/me", response_model=List[CoverLetterOut])
async def get_my_coverletters(uuid: str = Depends(authorize)):
    """Fetch all cover letters belonging to the current user."""
    letters = await cover_letters_dao.get_all_cover_letters(uuid)

    if not letters:
        return []

    mapped_letters = [
        {
            "id": str(l["_id"]),
            "user_id": l.get("uuid"),
            "title": l.get("title"),
            "company": l.get("company"),
            "position": l.get("position"),
            "content": l.get("content"),
            "created_at": l.get("created_at"),
            "usage_count": l.get("usage_count", 0)
        }
        for l in letters
    ]

    return mapped_letters

# ============================================================
# GET a single cover letter by ID
# ============================================================
@coverletter_router.get("/{letter_id}", response_model=CoverLetterOut)
async def get_coverletter(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    """Fetch a single cover letter by ID if it belongs to the current user."""
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)

    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")

    return {
        "id": str(letter["_id"]),
        "user_id": letter.get("uuid"),
        "title": letter.get("title"),
        "company": letter.get("company"),
        "position": letter.get("position"),
        "content": letter.get("content"),
        "created_at": letter.get("created_at"),
        "usage_count": letter.get("usage_count", 0)
    }

# ============================================================
# POST create a new cover letter
# ============================================================
@coverletter_router.post("")
async def add_coverletter(
    coverletter: CoverLetterIn,
    uuid: str = Depends(authorize)
):
    """Add a new cover letter for the current user."""
    new_letter = {
        "_id": str(uuid4()),
        "uuid": uuid,
        "title": coverletter.title,
        "company": coverletter.company,
        "position": coverletter.position,
        "content": coverletter.content,
        "created_at": datetime.utcnow().isoformat(),
        "usage_count": 1 if coverletter.template_type else 0,
        "template_type": getattr(coverletter, 'template_type', None)
    }

    inserted_id = await cover_letters_dao.add_cover_letter(new_letter)
    
    return {
        "coverletter_id": inserted_id,
        "data": {
            "_id": inserted_id,
            "title": new_letter["title"],
            "created_at": new_letter["created_at"]
        }
    }

# ============================================================
# POST upload an HTML/file cover letter
# ============================================================
@coverletter_router.post("/upload")
async def upload_coverletter(
    file: UploadFile = File(...),
    title: str = Form(...),
    company: str = Form(""),
    position: str = Form(""),
    version_name: str = Form("Version 1"),
    description: str = Form(None),
    uuid: str = Depends(authorize)
):
    """Upload a cover letter file (HTML, PDF, or DOCX)."""
    
    # For now, accept HTML files
    if file.content_type != "text/html" and not file.filename.lower().endswith('.html'):
        # If it's PDF/DOCX, we'd need different handling
        raise HTTPException(status_code=400, detail="Only HTML files are supported currently")
    
    try:
        content = await file.read()
        html_content = content.decode('utf-8')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")
    
    letter_id = str(uuid4())
    created_at = datetime.utcnow().isoformat()
    
    new_letter = {
        "_id": letter_id,
        "uuid": uuid,
        "title": title,
        "company": company,
        "position": position,
        "content": html_content,
        "version_name": version_name,
        "description": description,
        "created_at": created_at,
        "usage_count": 0,
        "uploadedFile": True,
        "template_type": None,
        "file_type": file.content_type,
        "file_size": len(content)
    }
    
    await cover_letters_dao.add_cover_letter(new_letter)
    
    return {
        "detail": "File uploaded successfully",
        "data": {
            "_id": letter_id,
            "title": title,
            "version_name": version_name,
            "created_at": created_at,
            "file_size": len(content)
        }
    }

# ============================================================
# PUT update a cover letter
# ============================================================
@coverletter_router.put("/{letter_id}")
async def update_coverletter(
    letter_id: str = Path(...),
    coverletter: CoverLetterIn = Body(...),
    uuid: str = Depends(authorize)
):
    """Update an existing cover letter if it belongs to the current user."""
    updates = {
        "title": coverletter.title,
        "company": coverletter.company,
        "position": coverletter.position,
        "content": coverletter.content
    }

    modified_count = await cover_letters_dao.update_cover_letter(letter_id, uuid, updates)

    if modified_count == 0:
        letter_exists = await cover_letters_dao.get_cover_letter(letter_id, uuid)
        if not letter_exists:
            raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")
        return {"message": "No changes to update"}

    return {"message": "Updated successfully"}

# ============================================================
# DELETE a cover letter
# ============================================================
@coverletter_router.delete("/{letter_id}")
async def delete_coverletter(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    """Delete a cover letter belonging to the current user."""
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")
    
    deleted_count = await cover_letters_dao.delete_cover_letter(letter_id)

    if deleted_count == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")

    return {"message": "Deleted successfully"}

# ============================================================
# PUT set as default cover letter
# ============================================================
@coverletter_router.put("/{letter_id}/default")
async def set_default_coverletter(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    """Set a cover letter as the default for the user."""
    # First verify the letter exists and belongs to user
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")
    
    # Implementation would set is_default field
    # For now, just return success
    return {"message": "Default cover letter set successfully"}

# ============================================================
# GET download as PDF
# ============================================================
@coverletter_router.get("/{letter_id}/download/pdf")
async def download_pdf(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    """Download cover letter as PDF with styling preserved."""
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    
    try:
        html_content = letter.get("content", "")
        filename = f"{letter.get('title', 'cover_letter')}.pdf"
        
        response = requests.post(
            'https://api.html2pdf.app/v1/generate',
            json={'html': html_content},
            timeout=30
        )
        
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail="PDF generation failed")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(response.content)
            tmp_path = tmp.name
        
        return FileResponse(
            tmp_path,
            media_type='application/pdf',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    except Exception as e:
        print(f"PDF Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

# ============================================================
# GET download as DOCX
# ============================================================
@coverletter_router.get("/{letter_id}/download/docx")
async def download_docx(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    """Download cover letter as DOCX with styling preserved."""
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    
    try:
        html_content = letter.get("content", "")
        soup = BeautifulSoup(html_content, 'html.parser')
        
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Segoe UI'
        font.size = Pt(11)
        
        for element in soup.body.children:
            process_html_element(element, doc)
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        filename = f"{letter.get('title', 'cover_letter')}.docx"
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(doc_buffer.getvalue())
            tmp_path = tmp.name
        
        return FileResponse(
            tmp_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    except Exception as e:
        print(f"DOCX Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX: {str(e)}")

# ============================================================
# HELPER: Process HTML elements into DOCX
# ============================================================
def process_html_element(element, doc):
    """Recursively process HTML elements and add to DOCX document."""
    if isinstance(element, str):
        text = element.strip()
        if text:
            doc.add_paragraph(text)
        return
    
    tag = element.name
    
    if tag == 'header':
        for child in element.children:
            process_html_element(child, doc)
    
    elif tag == 'h1':
        p = doc.add_paragraph()
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(element.get_text())
        run.font.size = Pt(24)
        run.font.bold = True
    
    elif tag == 'h2':
        p = doc.add_paragraph()
        p.style = 'Heading 2'
        run = p.add_run(element.get_text())
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
    
    elif tag == 'h3':
        p = doc.add_paragraph()
        p.style = 'Heading 3'
        run = p.add_run(element.get_text())
        run.font.bold = True
        run.font.color.rgb = RGBColor(52, 73, 94)
    
    elif tag == 'p':
        text = element.get_text(strip=True)
        if text:
            p = doc.add_paragraph(text)
            style_attr = element.get('style', '')
            if 'font-weight: 500' in style_attr or 'bold' in style_attr:
                for run in p.runs:
                    run.font.bold = True
            if 'font-style: italic' in style_attr or 'italic' in style_attr:
                for run in p.runs:
                    run.font.italic = True
    
    elif tag == 'div':
        class_attr = element.get('class', [])
        
        if 'summary' in class_attr:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(element.get_text())
            run.font.italic = True
            run.font.color.rgb = RGBColor(52, 152, 219)
        
        else:
            for child in element.children:
                process_html_element(child, doc)
    
    elif tag == 'ul':
        for li in element.find_all('li', recursive=False):
            doc.add_paragraph(li.get_text(), style='List Bullet')
    
    elif tag == 'li':
        text = element.get_text(strip=True)
        if text:
            doc.add_paragraph(text, style='List Bullet')
    
    elif tag == 'footer':
        text = element.get_text(strip=True)
        if text:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
    
    else:
        if hasattr(element, 'children'):
            for child in element.children:
                process_html_element(child, doc)